"""
PDF Converter - Convertidor de documentos
Autor: Carlos Ramz
Descripción: Script para convertir entre PDF, Word y Excel
"""

import PyPDF2
import os
from docx import Document
import pdfplumber

def leer_pdf_mejorado(ruta_archivo):
    """
    Lee un archivo PDF usando pdfplumber para texto normal o OCR para PDFs escaneados
    
    Args:
        ruta_archivo (str): Ruta al archivo PDF
        
    Returns:
        str: Texto extraído del PDF
    """
    try:
        import pdfplumber
        
        # Primero intentar extraer texto normal
        texto_completo = ""
        
        with pdfplumber.open(ruta_archivo) as pdf:
            print(f"El PDF tiene {len(pdf.pages)} páginas")
            
            for num_pagina, pagina in enumerate(pdf.pages, 1):
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
                print(f"Página {num_pagina} procesada")
        
        # Si no hay texto suficiente, usar OCR
        if len(texto_completo.strip()) < 50:
            print("🔍 Texto insuficiente detectado - Usando OCR...")
            texto_completo = extraer_texto_con_ocr(ruta_archivo)
        
        return texto_completo
        
    except ImportError:
        print("❌ Error: pdfplumber no está instalado")
        return None
    except FileNotFoundError:
        print(f"Error: No se encontró el archivo {ruta_archivo}")
        return None
    except Exception as e:
        print(f"Error al leer el PDF: {str(e)}")
        return None

def pdf_a_word(ruta_pdf, directorio_salida="./convertidos"):
    """
    Convierte un PDF a documento Word
    
    Args:
        ruta_pdf (str): Ruta al archivo PDF
        directorio_salida (str): Directorio donde guardar el archivo Word
        
    Returns:
        str: Ruta del archivo Word creado, o None si hay error
    """
    try:
        from docx import Document
        
        # Extraer texto del PDF
        texto_pdf = leer_pdf_mejorado(ruta_pdf)
        if not texto_pdf:
            print("❌ No se pudo extraer texto del PDF")
            return None
        
        # Crear directorio si no existe
        if not os.path.exists(directorio_salida):
            os.makedirs(directorio_salida)
            print(f"📁 Carpeta creada: {directorio_salida}")
        
        # Generar nombre del archivo Word
        nombre_archivo = os.path.basename(ruta_pdf)
        nombre_sin_extension = os.path.splitext(nombre_archivo)[0]
        nombre_word = f"{nombre_sin_extension}.docx"
        ruta_word = os.path.join(directorio_salida, nombre_word)
        
        # Crear documento Word
        documento = Document()
        
        # Agregar el texto manteniendo formato original
        # Dividir por líneas y mantener estructura exacta
        lineas = texto_pdf.split('\n')
        
        for linea in lineas:
            # Mantener líneas vacías para conservar espaciado
            if not linea.strip():
                documento.add_paragraph("")  # Línea vacía
            else:
                documento.add_paragraph(linea)  # Línea con contenido tal como está
        
        # Guardar el documento
        documento.save(ruta_word)
        
        print(f"✅ Archivo Word creado: {ruta_word}")
        return ruta_word
        
    except ImportError:
        print("❌ Error: python-docx no está instalado")
        return None
    except Exception as e:
        print(f"❌ Error al crear archivo Word: {str(e)}")
        return None

def pdf_a_excel(ruta_pdf, directorio_salida="./convertidos"):
    """
    Convierte un PDF a Excel (versión simple sin formato)
    """
    try:
        from openpyxl import Workbook
        
        # Extraer texto del PDF
        texto_pdf = leer_pdf_mejorado(ruta_pdf)
        if not texto_pdf:
            print("❌ No se pudo extraer texto del PDF")
            return None
        
        # Crear directorio si no existe
        if not os.path.exists(directorio_salida):
            os.makedirs(directorio_salida)
            print(f"📁 Carpeta creada: {directorio_salida}")
        
        # Generar nombre del archivo Excel
        nombre_archivo = os.path.basename(ruta_pdf)
        nombre_sin_extension = os.path.splitext(nombre_archivo)[0]
        nombre_excel = f"{nombre_sin_extension}.xlsx"
        ruta_excel = os.path.join(directorio_salida, nombre_excel)
        
        # Crear libro de Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Contenido PDF"
        
        # Encabezados simples
        ws['A1'] = "Tipo"
        ws['B1'] = "Contenido"
        ws['C1'] = "Nivel"
        
        # Procesar texto línea por línea
        lineas = texto_pdf.split('\n')
        fila_actual = 2
        
        for linea in lineas:
            linea_limpia = linea.strip()
            if not linea_limpia:
                continue
            
            tipo_contenido, nivel = detectar_tipo_contenido(linea_limpia)
            
            ws.cell(row=fila_actual, column=1, value=tipo_contenido)
            ws.cell(row=fila_actual, column=2, value=linea_limpia)
            ws.cell(row=fila_actual, column=3, value=nivel)
            
            fila_actual += 1
        
        # Ajustar ancho de columnas
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 8
        
        wb.save(ruta_excel)
        print(f"✅ Archivo Excel creado: {ruta_excel}")
        return ruta_excel
        
    except ImportError:
        print("❌ Error: openpyxl no está instalado")
        return None
    except Exception as e:
        print(f"❌ Error al crear archivo Excel: {str(e)}")
        return None
    
def detectar_tipo_contenido(linea):
    """
    Detecta el tipo de contenido de una línea
    
    Args:
        linea (str): Línea de texto a analizar
        
    Returns:
        tuple: (tipo_contenido, nivel)
    """
    # Título principal: mayúsculas, corto
    if linea.isupper() and len(linea) <= 50:
        return "TÍTULO", 1
    
    # Subtítulo: primera letra mayúscula, longitud media
    elif linea[0].isupper() and len(linea) <= 100 and not linea.endswith('.'):
        return "SUBTÍTULO", 2
    
    # Párrafo: texto normal, puede terminar en punto
    elif len(linea) > 20:
        return "PÁRRAFO", 3
    
    # Texto corto: elementos diversos
    else:
        return "TEXTO", 3

    
def detectar_pdf_escaneado(ruta_pdf):
    """
    Detecta si un PDF es escaneado (contiene solo imágenes)
    
    Args:
        ruta_pdf (str): Ruta al archivo PDF
        
    Returns:
        bool: True si es escaneado, False si tiene texto
    """
    try:
        import pdfplumber
        
        with pdfplumber.open(ruta_pdf) as pdf:
            # Revisar las primeras 3 páginas para detectar
            paginas_revisar = min(3, len(pdf.pages))
            
            for i in range(paginas_revisar):
                texto = pdf.pages[i].extract_text()
                if texto and len(texto.strip()) > 50:  # Si hay texto sustancial
                    return False
            
            return True  # No encontró texto = PDF escaneado
            
    except Exception as e:
        print(f"Error al detectar tipo de PDF: {str(e)}")
        return False

def extraer_texto_con_ocr(ruta_pdf):
    """
    Extrae texto de un PDF escaneado usando OCR
    
    Args:
        ruta_pdf (str): Ruta al archivo PDF
        
    Returns:
        str: Texto extraído con OCR
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
        import io
        
        # Configurar ruta de Tesseract (ajusta si es necesario)
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        texto_completo = ""
        
        # Abrir PDF
        doc = fitz.open(ruta_pdf)
        print(f"PDF escaneado detectado - {len(doc)} páginas para OCR")
        
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            
            # Convertir página a imagen
            mat = fitz.Matrix(2.0, 2.0)  # Zoom 2x para mejor calidad
            pix = pagina.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Procesar con OCR
            img = Image.open(io.BytesIO(img_data))
            texto_pagina = pytesseract.image_to_string(img, lang='eng')  # Inglés
            
            texto_completo += texto_pagina + "\n"
            print(f"Página {num_pagina + 1} procesada con OCR")
        
        doc.close()
        return texto_completo
        
    except ImportError as e:
        print(f"❌ Error: Falta instalar PyMuPDF: pip install PyMuPDF")
        return None
    except Exception as e:
        print(f"❌ Error en OCR: {str(e)}")
        return None


def mostrar_menu():
    """Muestra el menú principal del programa"""
    print("\n" + "="*50)
    print("🔄 PDF CONVERTER - MENÚ PRINCIPAL")
    print("="*50)
    print("1️⃣  Convertir PDF a Word")
    print("2️⃣  Convertir PDF a Excel") 
    print("3️⃣  Convertir PDF a Word y Excel")
    print("4️⃣  Listar archivos PDF disponibles")
    print("5️⃣  Salir del programa")
    print("="*50)

def listar_pdfs():
    """Lista todos los archivos PDF disponibles"""
    pdfs = [f for f in os.listdir('.') if f.endswith('.pdf')]
    
    if not pdfs:
        print("❌ No se encontraron archivos PDF en la carpeta actual")
        return []
    
    print("\n📁 Archivos PDF disponibles:")
    for i, pdf in enumerate(pdfs, 1):
        tamano = os.path.getsize(pdf)
        tamano_mb = tamano / (1024*1024)
        print(f"  {i}. {pdf} ({tamano_mb:.1f} MB)")
    
    return pdfs

def seleccionar_pdf():
    """Permite al usuario seleccionar un archivo PDF"""
    pdfs = listar_pdfs()
    
    if not pdfs:
        return None
    
    while True:
        try:
            opcion = input(f"\n📝 Selecciona un PDF (1-{len(pdfs)}) o 'q' para volver: ").strip()
            
            if opcion.lower() == 'q':
                return None
            
            indice = int(opcion) - 1
            if 0 <= indice < len(pdfs):
                return pdfs[indice]
            else:
                print(f"❌ Opción inválida. Ingresa un número entre 1 y {len(pdfs)}")
        
        except ValueError:
            print("❌ Por favor ingresa un número válido")

def procesar_conversion(pdf_seleccionado, tipo_conversion):
    """
    Procesa la conversión según el tipo seleccionado
    
    Args:
        pdf_seleccionado (str): Nombre del archivo PDF
        tipo_conversion (str): 'word', 'excel', 'ambos'
    """
    print(f"\n🔄 Procesando: {pdf_seleccionado}")
    print(f"📄 Tamaño: {os.path.getsize(pdf_seleccionado)} bytes")
    
    # Analizar el archivo primero
    print("\n🔍 Analizando archivo...")
    texto = leer_pdf_mejorado(pdf_seleccionado)
    
    if not texto:
        print("❌ No se pudo extraer texto del archivo")
        return
    
    # Mostrar muestra del texto
    print("\n📋 Muestra del texto extraído:")
    print("-" * 40)
    print(texto[:200] + "..." if len(texto) > 200 else texto)
    print("-" * 40)
    
    # Realizar conversiones
    if tipo_conversion in ['word', 'ambos']:
        print("\n🔄 Convirtiendo a Word...")
        archivo_word = pdf_a_word(pdf_seleccionado)
        if archivo_word:
            print(f"✅ Word creado: {archivo_word}")
        else:
            print("❌ Error al crear archivo Word")
    
    if tipo_conversion in ['excel', 'ambos']:
        print("\n🔄 Convirtiendo a Excel...")
        archivo_excel = pdf_a_excel(pdf_seleccionado)
        if archivo_excel:
            print(f"✅ Excel creado: {archivo_excel}")
        else:
            print("❌ Error al crear archivo Excel")
    
    print("\n✅ ¡Conversión completada!")

def main():
    """Función principal con menú interactivo"""
    print("🚀 PDF Converter v1.0.3 - Iniciando...")
    print("💡 Soporte para PDFs normales y escaneados (OCR)")
    
    while True:
        mostrar_menu()
        
        opcion = input("\n🎯 Selecciona una opción: ").strip()
        
        if opcion == '1':
            pdf = seleccionar_pdf()
            if pdf:
                procesar_conversion(pdf, 'word')
        
        elif opcion == '2':
            pdf = seleccionar_pdf()
            if pdf:
                procesar_conversion(pdf, 'excel')
        
        elif opcion == '3':
            pdf = seleccionar_pdf()
            if pdf:
                procesar_conversion(pdf, 'ambos')
        
        elif opcion == '4':
            listar_pdfs()
        
        elif opcion == '5':
            print("\n👋 ¡Gracias por usar PDF Converter!")
            print("🔗 Proyecto: https://github.com/carlosramz/pdf-converter")
            break
        
        else:
            print("❌ Opción inválida. Por favor selecciona 1-5")
        
        # Pausa antes de mostrar el menú de nuevo
        input("\n⏸️  Presiona Enter para continuar...")

if __name__ == "__main__":
    main()